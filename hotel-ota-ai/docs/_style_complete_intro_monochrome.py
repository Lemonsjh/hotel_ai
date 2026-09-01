from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SOURCE = Path(r"D:\酒店数字员工\hotel--ota-ai\docs\酒店OTA数字员工与数据采集执行服务项目完整介绍.docx")
OUTPUT = Path(r"D:\酒店数字员工\hotel--ota-ai\docs\酒店OTA数字员工与数据采集执行服务项目完整介绍-黑白排版版.docx")

BLACK = "000000"
DARK_GRAY = "3A3A3A"
MID_GRAY = "737373"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"


def remove_children(parent, tag):
    for child in list(parent):
        if child.tag == qn(tag):
            parent.remove(child)


def set_font(run, size=None, bold=None, color=BLACK, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    remove_children(tc_pr, "w:shd")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="BFBFBF", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_cell_margins(cell, top=105, start=130, bottom=105, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    table_width = sum(widths)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def add_paragraph_rule(paragraph, color="9B9B9B", size="6", space="7"):
    p_pr = paragraph._p.get_or_add_pPr()
    remove_children(p_pr, "w:pBdr")
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, before, after in (("Heading 1", 15, 18, 8), ("Heading 2", 12, 13, 6), ("Heading 3", 10.5, 10, 4)):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def style_title_block(document):
    nonempty = [p for p in document.paragraphs if p.text.strip()]
    if not nonempty:
        return
    title = nonempty[0]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(7)
    for run in title.runs:
        set_font(run, size=24, bold=True, color=BLACK)
    add_paragraph_rule(title, color="000000", size="10", space="9")
    if len(nonempty) > 1:
        subtitle = nonempty[1]
        subtitle.paragraph_format.space_after = Pt(16)
        for run in subtitle.runs:
            set_font(run, size=12, bold=False, color=MID_GRAY)


def style_paragraphs(document):
    part_titles = {"第二部分  酒店数据采集与 OTA 执行服务", "第三部分  两个系统的端到端协同"}
    callout_prefixes = ("一句话概括：", "任务表是", "关键边界：", "能力适用说明：", "本部分完整介绍")
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if text in part_titles:
            paragraph.style = document.styles["Heading 1"]
            paragraph.paragraph_format.page_break_before = True
        if paragraph.style and paragraph.style.name == "Heading 1":
            paragraph.paragraph_format.keep_with_next = True
            add_paragraph_rule(paragraph, color="000000", size="8", space="5")
            for run in paragraph.runs:
                set_font(run, size=15, bold=True, color=BLACK)
        elif paragraph.style and paragraph.style.name == "Heading 2":
            paragraph.paragraph_format.keep_with_next = True
            add_paragraph_rule(paragraph, color="A6A6A6", size="4", space="3")
            for run in paragraph.runs:
                set_font(run, size=12, bold=True, color=BLACK)
        elif text.startswith(callout_prefixes):
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(8)
            for run in paragraph.runs:
                set_font(run, size=10.5, bold=False, color=DARK_GRAY)
        else:
            for run in paragraph.runs:
                if run.text.strip():
                    set_font(run, size=10.5, color=DARK_GRAY)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    col_count = len(table.columns)
    usable = 9360
    if col_count == 2:
        widths = [2600, 6760]
    elif col_count == 3:
        widths = [1800, 4000, 3560]
    elif col_count == 4:
        widths = [1500, 2750, 2750, 2360]
    else:
        base = usable // col_count
        widths = [base] * col_count
        widths[-1] += usable - sum(widths)
    set_table_geometry(table, widths)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell, color="A6A6A6" if row_idx == 0 else "D0D0D0", size="5" if row_idx == 0 else "3")
            set_cell_shading(cell, BLACK if row_idx == 0 else (LIGHT_GRAY if row_idx % 2 == 0 else WHITE))
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_font(run, size=9.2, bold=(row_idx == 0), color=WHITE if row_idx == 0 else DARK_GRAY)


def style_headers_footers(document):
    for section in document.sections:
        section.top_margin = Cm(2.25)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)
        for paragraph in section.header.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.space_after = Pt(2)
            add_paragraph_rule(paragraph, color="BFBFBF", size="3", space="2")
            for run in paragraph.runs:
                set_font(run, size=8, color=MID_GRAY)
        for paragraph in section.footer.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(2)
            for run in paragraph.runs:
                set_font(run, size=8, color=MID_GRAY)


def main():
    document = Document(SOURCE)
    configure_styles(document)
    style_title_block(document)
    style_paragraphs(document)
    for table in document.tables:
        style_table(table)
    style_headers_footers(document)
    document.core_properties.title = "酒店 OTA 数字员工与数据采集执行服务项目完整介绍"
    document.core_properties.subject = "黑白商务排版版"
    document.save(OUTPUT)
    print(f"saved={OUTPUT}")
    print(f"paragraphs={len(document.paragraphs)} tables={len(document.tables)}")


if __name__ == "__main__":
    main()
